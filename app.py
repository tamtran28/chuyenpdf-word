import streamlit as st
import fitz
import os
import subprocess
import shutil
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt

# =======================
# Utility
# =======================

def has_extractable_text(pdf_path: str, sample_pages: int = 3, char_threshold: int = 40) -> bool:
    doc = fitz.open(pdf_path)
    n = min(sample_pages, len(doc))
    total = 0
    for i in range(n):
        total += len(doc[i].get_text("text") or "")
    doc.close()
    return total >= char_threshold


def check_ocr_tools():
    for exe in ("tesseract", "qpdf", "gs"):
        if shutil.which(exe) is None:
            return False, f"Thiếu tool: {exe}"
    return True, "OK"


def ocr_pdf(input_pdf: str, output_pdf: str, lang="vie"):
    cmd = [
        "ocrmypdf",
        "--force-ocr",
        "--rotate-pages",
        "--language", lang,
        "--output-type", "pdf",
        input_pdf,
        output_pdf
    ]
    subprocess.run(cmd, check=True)


def pdf_to_docx_safe(input_pdf, output_docx):
    with fitz.open(input_pdf) as d:
        last = max(0, d.page_count - 1)
    cv = Converter(input_pdf)
    try:
        cv.convert(output_docx, start=0, end=last)
    finally:
        cv.close()


def raster_to_docx(input_pdf, output_docx, dpi=200):
    doc = fitz.open(input_pdf)
    word = Document()
    for i, page in enumerate(doc, start=1):
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        img_path = f"tmp_page_{i}.png"
        pix.save(img_path)
        word.add_picture(img_path, width=Inches(6.5))
        os.remove(img_path)
    word.save(output_docx)


def normalize_docx_font(docx_path, font_name="Times New Roman", font_size=12):
    try:
        doc = Document(docx_path)
        for p in doc.paragraphs:
            for r in p.runs:
                r.font.name = font_name
                r.font.size = Pt(font_size)

        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = font_name
                            r.font.size = Pt(font_size)

        doc.save(docx_path)
    except Exception as e:
        st.warning(f"Không thể chuẩn hoá font: {e}")


# =======================
# Streamlit UI
# =======================

st.title("📄 PDF → DOCX Converter (có OCR)")
st.write("Ứng dụng Streamlit từ mã gốc Colab của bạn.")

uploaded_files = st.file_uploader("Tải lên file PDF", type=["pdf"], accept_multiple_files=True)

APPLY_FONT_NORMALIZE = st.checkbox("Chuẩn hoá font Times New Roman 12", value=True)
FORCE_OCR = st.checkbox("Bắt buộc OCR (ngay cả khi PDF có text)")

if uploaded_files:
    OCR_OK, msg = check_ocr_tools()
    if not OCR_OK:
        st.warning(f"⚠️ OCR không khả dụng: {msg}")

    for file in uploaded_files:
        pdf_path = f"input_{file.name}"
        with open(pdf_path, "wb") as f:
            f.write(file.read())

        base = os.path.splitext(file.name)[0]
        ocr_pdf_path = f"{base}.ocr.pdf"
        output_docx = f"{base}.docx"

        need_ocr = FORCE_OCR or not has_extractable_text(pdf_path)
        src = pdf_path

        st.write(f"### 🔧 Xử lý file: **{file.name}**")
        st.write(f"- OCR cần thiết: {need_ocr}")

        if need_ocr and OCR_OK:
            try:
                st.write("🔍 Đang OCR...")
                ocr_pdf(pdf_path, ocr_pdf_path)
                src = ocr_pdf_path
                st.success("OCR thành công.")
            except Exception as e:
                st.error(f"OCR lỗi: {e}")

        with st.spinner("Đang chuyển PDF → DOCX..."):
            try:
                pdf_to_docx_safe(src, output_docx)
                if APPLY_FONT_NORMALIZE:
                    normalize_docx_font(output_docx)
                st.success("✔ Chuyển đổi PDF → DOCX thành công!")
            except Exception as e:
                st.error(f"Lỗi pdf2docx: {e}")
                st.info("Thử fallback bằng raster ảnh.")
                raster_to_docx(src, output_docx)
                st.success("✔ Fallback ảnh → DOCX thành công!")

        with open(output_docx, "rb") as f:
            st.download_button(
                label=f"⬇ Tải xuống {output_docx}",
                data=f,
                file_name=output_docx
            )
